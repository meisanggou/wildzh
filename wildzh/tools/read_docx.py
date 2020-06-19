# !/usr/bin/env python
# coding: utf-8
import os
import pdb
import re
from win32com import client as wc
import docx
from docx.enum.style import WD_STYLE_TYPE


__author__ = 'meisa'


def doc_to_docx(doc_path):
    word = wc.Dispatch("Word.Application")
    doc = word.Documents.Open(doc_path)
    docx_path = doc_path + "x"
    doc.SaveAs(docx_path, 12)  # 12为docx
    doc.Close()
    word.Quit()
    return docx_path


Q_TYPE_COMP = re.compile(u"(选择|名词解释|简答题|计算题|论述题)")


def get_question(question_items, *args):
    if len(question_items) == 0:
        return None
    s_c = args
    desc = ""
    q_no = question_items[0]
    i = 1
    while i < len(question_items):
        qs_item = question_items[i]
        if re.match("A|B|C|D", qs_item, re.I) is not None:
            break
        desc += qs_item
        i += 1
    options = []

    def replace(match_r):
        mgs = match_r.groups()
        return "\n"
    # print(u"(A|B|C|D)(%s)" % "|".join(map(lambda x: re.escape(x), s_c)))
    r_compile = re.compile(u"(^|\s)(A|B|C|D)(%s)" % "|".join(map(lambda x: re.escape(x), s_c)))
    for item in question_items[i:]:
        r_item = r_compile.sub(replace, item)
        options.extend(r_item.split("\n"))
    real_options = map(lambda x: x.strip(), options)
    real_options = filter(lambda x: len(x) > 0, real_options)
    if len(real_options) != 4:
        for r_o in real_options:
            print(r_o)
        print(q_no)
        print(question_items)
        return None
    # for r_o in real_options:
    #     print(r_o)
    return dict(no=q_no, desc=desc, options=real_options)


def find_member(file_path, *args):
    s_c = args
    m_compile = re.compile(u"(\d+)(%s)" % "|".join(map(lambda x: re.escape(x), s_c)))
    doc = docx.Document(file_path)
    questions_s = []
    current_q_type = None
    current_question = []
    title = doc.paragraphs[0].text
    print(title)
    for pg in doc.paragraphs[1:]:
        s_line = pg.text.strip()
        if len(s_line) <= 0:
            continue
        # 判断是否是题目类型
        type_mr = Q_TYPE_COMP.findall(s_line)
        # print(pg.text)
        if len(type_mr) == 1:
            current_q_type = type_mr[0]
            continue
        # 判断是否是题目
        m_no = m_compile.match(s_line)
        if m_no is not None:
            q_item = get_question(current_question, ".", u"、", u"．")
            if q_item is not None:
                questions_s.append(q_item)
            q_no = int(m_no.groups()[0])
            current_question = [q_no]
            current_question.append(s_line[len("".join(m_no.groups())):])
            if q_no == 20:
                pdb.set_trace()
        else:
            current_question.append(s_line)
        # 判断是否是选项
        # print(pg.text)
    return []


def find_from_dir(directory_name):
    files = os.listdir(directory_name)
    all_member = []
    for file_item in files:
        if file_item.startswith("~$"):
            continue
        file_path = os.path.join(directory_name, file_item)
        items = os.path.split(file_path)

        if os.path.isfile(file_path) is False:
            continue
        elif file_path.endswith("答案.docx") is True:
            continue
        elif file_path.endswith(".doc") is True:
            if os.path.exists(file_path + "x"):
                continue
            file_path = doc_to_docx(file_path)
        elif file_path.endswith(".docx") is False:
            print(u"跳过文件 %s" % file_path)
            continue
        members = find_member(file_path, ".", u"、", u"．")
        if len(members) <= 0:
            print(u"请检查文件%s" % file_path)
        all_member.extend(members)
    # for mem in all_member:
    #     print("\t".join(mem))
    all_member.sort(key=lambda x: x[3])
    return all_member


all_members = find_from_dir(r'D:\Project\word')
print(all_members)
